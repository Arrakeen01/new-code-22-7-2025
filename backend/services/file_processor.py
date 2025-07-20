import os
import io
import tempfile
import zipfile
import aiofiles
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from PyPDF2 import PdfReader
from docx import Document
import magic
import base64

from models import UploadedFile, FileType


class FileProcessor:
    """Service for processing uploaded files and extracting content"""
    
    SUPPORTED_CODE_EXTENSIONS = {
        '.py', '.js', '.jsx', '.ts', '.tsx', '.java', '.cpp', '.c', '.h', 
        '.php', '.rb', '.go', '.rs', '.swift', '.kt', '.scala', '.cs',
        '.html', '.css', '.scss', '.sass', '.vue', '.svelte', '.json',
        '.xml', '.yaml', '.yml', '.md', '.txt', '.sql', '.sh', '.bat'
    }
    
    SUPPORTED_SRS_EXTENSIONS = {
        '.pdf', '.doc', '.docx', '.md', '.txt'
    }
    
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    MAX_CODE_FILES = 100
    MAX_SRS_FILES = 10

    @staticmethod
    def validate_file(file_name: str, file_size: int, file_type: FileType) -> Tuple[bool, str]:
        """Validate uploaded file"""
        # Check file size
        if file_size > FileProcessor.MAX_FILE_SIZE:
            return False, f"File size exceeds {FileProcessor.MAX_FILE_SIZE / (1024*1024)}MB limit"
        
        # Get file extension
        file_extension = Path(file_name).suffix.lower()
        
        # Check file type and extension
        if file_type == FileType.CODE:
            if file_extension not in FileProcessor.SUPPORTED_CODE_EXTENSIONS:
                return False, f"Unsupported code file extension: {file_extension}"
        elif file_type == FileType.SRS:
            if file_extension not in FileProcessor.SUPPORTED_SRS_EXTENSIONS:
                return False, f"Unsupported SRS file extension: {file_extension}"
        
        return True, "Valid file"

    @staticmethod
    async def process_uploaded_file(
        file_name: str, 
        file_content: str, 
        file_type: FileType,
        mime_type: str,
        session_id: str
    ) -> UploadedFile:
        """Process a single uploaded file"""
        
        # Decode base64 content if needed
        try:
            if file_content.startswith('data:'):
                # Remove data URL prefix
                file_content = file_content.split(',')[1]
            
            # For text files, try to decode as text
            if file_type == FileType.CODE or mime_type.startswith('text/'):
                try:
                    decoded_content = base64.b64decode(file_content).decode('utf-8')
                except:
                    decoded_content = file_content
            else:
                decoded_content = file_content
                
        except Exception as e:
            print(f"Error decoding file content: {e}")
            decoded_content = file_content

        uploaded_file = UploadedFile(
            name=file_name,
            type=file_type,
            size=len(file_content),
            content=decoded_content,
            mime_type=mime_type,
            session_id=session_id
        )
        
        return uploaded_file

    @staticmethod
    async def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_stream = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_stream)
            
            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            
            return text_content.strip()
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return f"Error reading PDF: {str(e)}"

    @staticmethod
    async def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_stream = io.BytesIO(file_content)
            doc = Document(docx_stream)
            
            text_content = []
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            return '\n'.join(text_content)
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
            return f"Error reading DOCX: {str(e)}"

    @staticmethod
    async def process_srs_file(uploaded_file: UploadedFile) -> str:
        """Process SRS file and extract text content"""
        file_extension = Path(uploaded_file.name).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                # Decode base64 content to bytes
                file_bytes = base64.b64decode(uploaded_file.content)
                return await FileProcessor.extract_text_from_pdf(file_bytes)
            
            elif file_extension in ['.doc', '.docx']:
                # Decode base64 content to bytes
                file_bytes = base64.b64decode(uploaded_file.content)
                return await FileProcessor.extract_text_from_docx(file_bytes)
            
            elif file_extension in ['.md', '.txt']:
                # For text files, content should already be decoded
                return uploaded_file.content
            
            else:
                return f"Unsupported file type: {file_extension}"
                
        except Exception as e:
            print(f"Error processing SRS file {uploaded_file.name}: {e}")
            return f"Error processing {uploaded_file.name}: {str(e)}"

    @staticmethod
    def detect_programming_language(file_name: str, content: str) -> str:
        """Detect programming language from file extension and content"""
        file_extension = Path(file_name).suffix.lower()
        
        language_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.jsx': 'javascript',
            '.ts': 'typescript',
            '.tsx': 'typescript',
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c',
            '.h': 'c',
            '.php': 'php',
            '.rb': 'ruby',
            '.go': 'go',
            '.rs': 'rust',
            '.swift': 'swift',
            '.kt': 'kotlin',
            '.scala': 'scala',
            '.cs': 'csharp',
            '.html': 'html',
            '.css': 'css',
            '.scss': 'scss',
            '.sass': 'sass',
            '.vue': 'vue',
            '.svelte': 'svelte',
            '.json': 'json',
            '.xml': 'xml',
            '.yaml': 'yaml',
            '.yml': 'yaml',
            '.md': 'markdown',
            '.sql': 'sql',
            '.sh': 'bash',
            '.bat': 'batch'
        }
        
        return language_map.get(file_extension, 'text')

    @staticmethod
    async def process_zip_file(file_content: bytes, session_id: str) -> List[UploadedFile]:
        """Process ZIP file and extract individual files"""
        processed_files = []
        
        try:
            zip_stream = io.BytesIO(file_content)
            with zipfile.ZipFile(zip_stream, 'r') as zip_file:
                for file_info in zip_file.filelist:
                    if file_info.is_dir():
                        continue
                    
                    # Skip hidden files and system files
                    if file_info.filename.startswith('.') or '__MACOSX' in file_info.filename:
                        continue
                    
                    # Check if it's a supported code file
                    file_extension = Path(file_info.filename).suffix.lower()
                    if file_extension not in FileProcessor.SUPPORTED_CODE_EXTENSIONS:
                        continue
                    
                    # Extract file content
                    with zip_file.open(file_info) as extracted_file:
                        file_content = extracted_file.read()
                        
                        # Try to decode as text
                        try:
                            text_content = file_content.decode('utf-8')
                        except UnicodeDecodeError:
                            try:
                                text_content = file_content.decode('latin-1')
                            except:
                                # Skip binary files
                                continue
                        
                        uploaded_file = UploadedFile(
                            name=file_info.filename,
                            type=FileType.CODE,
                            size=len(text_content),
                            content=text_content,
                            mime_type='text/plain',
                            session_id=session_id
                        )
                        processed_files.append(uploaded_file)
                        
                        # Limit number of files
                        if len(processed_files) >= FileProcessor.MAX_CODE_FILES:
                            break
        
        except Exception as e:
            print(f"Error processing ZIP file: {e}")
            raise Exception(f"Error processing ZIP file: {str(e)}")
        
        return processed_files

    @staticmethod
    def get_file_stats(files: List[UploadedFile]) -> Dict:
        """Get statistics about uploaded files"""
        stats = {
            'total_files': len(files),
            'code_files': len([f for f in files if f.type == FileType.CODE]),
            'srs_files': len([f for f in files if f.type == FileType.SRS]),
            'total_size': sum(f.size for f in files),
            'languages': {}
        }
        
        # Count languages
        for file in files:
            if file.type == FileType.CODE:
                lang = FileProcessor.detect_programming_language(file.name, file.content or "")
                stats['languages'][lang] = stats['languages'].get(lang, 0) + 1
        
        return stats